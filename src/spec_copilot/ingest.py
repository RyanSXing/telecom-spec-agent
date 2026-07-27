import hashlib
import io
import re
import zipfile
from collections.abc import Iterator, Sequence
from pathlib import Path, PurePosixPath
from typing import Any
from urllib.parse import urlparse

import httpx
from docx import Document as open_docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from llama_index.core.node_parser import NodeParser, SentenceSplitter
from llama_index.core.schema import BaseNode, Document, TextNode
from pydantic import Field

from spec_copilot.models import Chunk, Clause

_HEADING_STYLE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)
_NUMBERED_HEADING = re.compile(
    r"^\s*([0-9]+(?:\.[0-9A-Za-z]+)*|[A-Z](?:\.[0-9A-Za-z]+)+)"
    r"(?:\t+|\s{2,})(.+?)\s*$"
)
_ANNEX_HEADING = re.compile(
    r"^\s*Annex\s+([A-Z])(?:\s+\([^)]+\))?\s*:?\s*(?:\t+)?(.*?)\s*$",
    re.IGNORECASE,
)
SPEC_URL = "https://www.3gpp.org/ftp/Specs/archive/23_series/23.501/23501-hf0.zip"
SPEC_SHA256 = "26a44ebac62fb954d8be7747eaade48c9bf5949867a078bf11c6a445fc0b7ace"
SPEC_MEMBER = "23501-hf0.docx"


def download_docx_archive(
    url: str,
    expected_sha256: str,
    member_name: str,
    destination: Path,
    client: httpx.Client | None = None,
) -> Path:
    destination.mkdir(parents=True, exist_ok=True)
    output_path = destination / Path(member_name).name
    archive_name = PurePosixPath(urlparse(url).path).name or "specification.zip"
    archive_path = destination / archive_name
    payload = archive_path.read_bytes() if archive_path.exists() else b""
    downloaded = False

    if hashlib.sha256(payload).hexdigest() != expected_sha256:
        owns_client = client is None
        http = client or httpx.Client(
            follow_redirects=True,
            timeout=120,
            headers={"User-Agent": "Spec-Copilot/0.1 (+public 3GPP research demo)"},
        )
        try:
            response = http.get(url)
            response.raise_for_status()
            payload = response.content
        finally:
            if owns_client:
                http.close()

        if hashlib.sha256(payload).hexdigest() != expected_sha256:
            raise ValueError("downloaded archive hash does not match the pinned specification")
        downloaded = True

    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        for info in archive.infolist():
            path = PurePosixPath(info.filename)
            if path.is_absolute() or ".." in path.parts:
                raise ValueError(f"unsafe ZIP member: {info.filename}")
        if member_name not in archive.namelist():
            raise ValueError(f"missing ZIP member: {member_name}")
        content = archive.read(member_name)

    if downloaded:
        temporary_archive = archive_path.with_suffix(f"{archive_path.suffix}.part")
        temporary_archive.write_bytes(payload)
        temporary_archive.replace(archive_path)
    temporary_path = output_path.with_suffix(f"{output_path.suffix}.part")
    temporary_path.write_bytes(content)
    temporary_path.replace(output_path)
    return output_path


def fetch_pinned_spec(destination: Path = Path("data/raw")) -> Path:
    return download_docx_archive(
        SPEC_URL,
        SPEC_SHA256,
        SPEC_MEMBER,
        destination,
    )


def _blocks(document: DocxDocument) -> Iterator[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _heading(
    paragraph: Paragraph,
    *,
    allow_unstyled: bool = True,
) -> tuple[int, str, str] | None:
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.upper().startswith("TOC"):
        return None

    style_match = _HEADING_STYLE.match(style_name)
    if style_match is None and not allow_unstyled:
        return None
    text = paragraph.text.strip()
    annex_match = _ANNEX_HEADING.match(text)
    if annex_match:
        number, title = annex_match.groups()
        level = int(style_match.group(1)) if style_match else 1
        return level, number.upper(), title.strip() or f"Annex {number.upper()}"

    heading_match = _NUMBERED_HEADING.match(text)
    if heading_match:
        number, title = heading_match.groups()
        level = int(style_match.group(1)) if style_match else number.count(".") + 1
        return level, number, title.strip()
    return None


def _table_text(table: Table) -> str:
    rows = [[" ".join(cell.text.split()) for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    header = "| " + " | ".join(rows[0]) + " |"
    divider = "| " + " | ".join("---" for _ in range(width)) + " |"
    body = ["| " + " | ".join(row) + " |" for row in rows[1:]]
    return "\n".join([header, divider, *body])


def parse_docx(
    path: Path,
    spec_id: str,
    release: str,
    source_url: str,
) -> list[Clause]:
    document = open_docx(path)
    has_heading_styles = any(
        _HEADING_STYLE.match(paragraph.style.name if paragraph.style else "")
        for paragraph in document.paragraphs
    )
    clauses: list[Clause] = []
    current: Clause | None = None
    parents: list[tuple[int, str]] = []

    for block in _blocks(document):
        if isinstance(block, Paragraph):
            heading = _heading(block, allow_unstyled=not has_heading_styles)
            if heading:
                level, number, title = heading
                while parents and parents[-1][0] >= level:
                    parents.pop()
                path_numbers = [parent_number for _, parent_number in parents] + [number]
                current = Clause(
                    spec_id=spec_id,
                    release=release,
                    clause_number=number,
                    title=title,
                    clause_path=path_numbers,
                    source_url=source_url,
                )
                clauses.append(current)
                parents.append((level, number))
                continue

            text = block.text.strip()
            if current is not None and text and not block.style.name.upper().startswith("TOC"):
                current.text = f"{current.text}\n\n{text}".strip()
        elif current is not None:
            table_text = _table_text(block)
            if table_text:
                current.tables.append(table_text)

    return clauses


class ClauseNodeParser(NodeParser):
    chunk_size: int = Field(default=800, gt=0)

    def _parse_nodes(
        self,
        nodes: Sequence[BaseNode],
        show_progress: bool = False,
        **kwargs: Any,
    ) -> list[BaseNode]:
        del show_progress, kwargs
        splitter = SentenceSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=0,
            include_metadata=False,
        )
        parsed: list[BaseNode] = []
        for node in nodes:
            header = str(node.metadata["header"])
            for chunk_index, piece in enumerate(splitter.split_text(node.text)):
                text = f"{header}\n\n{piece.strip()}"
                metadata = {**node.metadata, "chunk_index": chunk_index}
                parsed.append(
                    TextNode(
                        id_=hashlib.sha256(text.encode()).hexdigest(),
                        text=text,
                        metadata=metadata,
                    )
                )
        return parsed

    def chunk_clauses(self, clauses: Sequence[Clause]) -> list[Chunk]:
        documents = []
        for clause in clauses:
            body = "\n\n".join(part for part in [clause.text, *clause.tables] if part)
            if not body:
                continue
            header = f"[{clause.spec_id} §{clause.clause_number} {clause.title}]"
            documents.append(
                Document(
                    text=body,
                    metadata={
                        "header": header,
                        "spec_id": clause.spec_id,
                        "release": clause.release,
                        "clause_number": clause.clause_number,
                        "title": clause.title,
                        "clause_path": clause.clause_path,
                        "source_url": clause.source_url,
                    },
                )
            )

        return [
            Chunk(
                chunk_id=node.node_id,
                chunk_index=int(node.metadata["chunk_index"]),
                spec_id=str(node.metadata["spec_id"]),
                release=str(node.metadata["release"]),
                clause_number=str(node.metadata["clause_number"]),
                title=str(node.metadata["title"]),
                clause_path=list(node.metadata["clause_path"]),
                text=node.text,
                source_url=str(node.metadata["source_url"]),
            )
            for node in self.get_nodes_from_documents(documents)
        ]
