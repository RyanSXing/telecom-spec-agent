import hashlib
import io
import zipfile
from pathlib import Path

import httpx
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from spec_copilot.ingest import ClauseNodeParser, download_docx_archive, parse_docx


def make_spec(path: Path) -> None:
    document = Document()
    document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("1\tFake contents entry", style="TOC 1")
    document.add_paragraph("This must not be indexed.")

    document.add_paragraph("1\tScope", style="Heading 1")
    document.add_paragraph("Scope paragraph.")
    document.add_paragraph("9\tThis is a numbered body item.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Function"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "AMF"
    table.cell(1, 1).text = "Access"

    document.add_paragraph("1.1\tDetails", style="Heading 2")
    document.add_paragraph("First detail paragraph. " * 20)
    document.add_paragraph("Second detail paragraph. " * 20)
    document.add_paragraph("2\tNext clause", style="Heading 1")
    document.add_paragraph("Next clause body.")
    document.add_paragraph("Annex A (informative):\tExamples", style="Heading 1")
    document.add_paragraph("Annex introduction.")
    document.add_paragraph("A.1\tExample case", style="Heading 2")
    document.add_paragraph("Example body.")
    document.save(path)


def test_parse_docx_preserves_clause_structure_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "fixture.docx"
    make_spec(path)

    clauses = parse_docx(
        path,
        spec_id="TS 23.501",
        release="Rel-17",
        source_url="https://example.test/spec.zip",
    )

    assert [clause.clause_number for clause in clauses] == ["1", "1.1", "2", "A", "A.1"]
    assert clauses[0].title == "Scope"
    assert clauses[0].clause_path == ["1"]
    assert clauses[0].tables == ["| Function | Role |\n| --- | --- |\n| AMF | Access |"]
    assert clauses[1].clause_path == ["1", "1.1"]
    assert clauses[3].title == "Examples"
    assert clauses[4].clause_path == ["A", "A.1"]
    assert "numbered body item" in clauses[0].text
    assert all("Fake contents" not in clause.text for clause in clauses)


def test_parse_docx_falls_back_to_numbered_unstyled_headings(tmp_path: Path) -> None:
    path = tmp_path / "fallback.docx"
    document = Document()
    document.add_paragraph("3\tFallback clause")
    document.add_paragraph("Fallback body.")
    document.add_paragraph("3.1\tNested fallback")
    document.add_paragraph("Nested body.")
    document.save(path)

    clauses = parse_docx(path, "TS 23.501", "Rel-17", "https://example.test/spec.zip")

    assert [clause.clause_number for clause in clauses] == ["3", "3.1"]
    assert clauses[1].clause_path == ["3", "3.1"]


def test_clause_node_parser_never_mixes_clauses(tmp_path: Path) -> None:
    path = tmp_path / "fixture.docx"
    make_spec(path)
    clauses = parse_docx(path, "TS 23.501", "Rel-17", "https://example.test/spec.zip")

    chunks = ClauseNodeParser(chunk_size=18).chunk_clauses(clauses)

    detail_chunks = [chunk for chunk in chunks if chunk.clause_number == "1.1"]
    assert len(detail_chunks) >= 2
    assert all(chunk.text.startswith("[TS 23.501 §1.1 Details]") for chunk in detail_chunks)
    assert all("Next clause body." not in chunk.text for chunk in detail_chunks)
    assert all(
        chunk.chunk_id == hashlib.sha256(chunk.text.encode()).hexdigest() for chunk in chunks
    )


def zip_bytes(member_name: str, content: bytes, unsafe: bool = False) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr(member_name, content)
        if unsafe:
            archive.writestr("../escape.txt", "no")
    return buffer.getvalue()


def test_download_docx_archive_verifies_hash_and_extracts_one_member(tmp_path: Path) -> None:
    payload = zip_bytes("spec.docx", b"docx bytes")
    client = httpx.Client(
        transport=httpx.MockTransport(lambda request: httpx.Response(200, content=payload))
    )

    path = download_docx_archive(
        "https://example.test/spec.zip",
        hashlib.sha256(payload).hexdigest(),
        "spec.docx",
        tmp_path,
        client,
    )

    assert path.read_bytes() == b"docx bytes"


def test_download_docx_archive_rejects_unsafe_members(tmp_path: Path) -> None:
    payload = zip_bytes("spec.docx", b"docx bytes", unsafe=True)
    client = httpx.Client(
        transport=httpx.MockTransport(lambda request: httpx.Response(200, content=payload))
    )

    with pytest.raises(ValueError, match="unsafe ZIP member"):
        download_docx_archive(
            "https://example.test/spec.zip",
            hashlib.sha256(payload).hexdigest(),
            "spec.docx",
            tmp_path,
            client,
        )
    assert not (tmp_path / "spec.zip").exists()


def test_download_docx_archive_reuses_verified_cached_archive(tmp_path: Path) -> None:
    payload = zip_bytes("spec.docx", b"docx bytes")
    requests = 0

    def respond(request: httpx.Request) -> httpx.Response:
        nonlocal requests
        requests += 1
        return httpx.Response(200, content=payload)

    client = httpx.Client(transport=httpx.MockTransport(respond))
    arguments = (
        "https://example.test/spec.zip",
        hashlib.sha256(payload).hexdigest(),
        "spec.docx",
        tmp_path,
        client,
    )

    download_docx_archive(*arguments)
    download_docx_archive(*arguments)

    assert requests == 1
