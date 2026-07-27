import json
from pathlib import Path

from docx import Document

from spec_copilot import cli
from spec_copilot.pipeline import ingest_path


class FakeEmbedder:
    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [[float(index), 0.0] for index, _ in enumerate(texts)]


class FakeStore:
    def __init__(self) -> None:
        self.ensured = False
        self.chunks = []
        self.vectors = []

    def ensure_index(self) -> None:
        self.ensured = True

    def index_chunks(self, chunks, vectors) -> None:
        self.chunks = chunks
        self.vectors = vectors


def test_ingest_path_parses_embeds_and_indexes(tmp_path: Path) -> None:
    path = tmp_path / "spec.docx"
    document = Document()
    document.add_paragraph("1\tScope", style="Heading 1")
    document.add_paragraph("Architecture scope.")
    document.save(path)
    store = FakeStore()

    summary = ingest_path(
        path,
        "https://example.test/spec.zip",
        FakeEmbedder(),
        store,
    )

    assert store.ensured is True
    assert store.chunks[0].clause_number == "1"
    assert len(store.chunks) == len(store.vectors) == 1
    assert summary["documents"] == 1
    assert summary["clauses"] == 1
    assert summary["chunks"] == 1


def test_cli_routes_ingest_and_eval_commands(monkeypatch, capsys) -> None:
    monkeypatch.setattr(cli, "run_ingest", lambda settings: {"chunks": 7})
    monkeypatch.setattr(cli, "run_evaluation", lambda settings: {"mrr": 0.5})

    assert cli.main(["ingest"]) == 0
    assert json.loads(capsys.readouterr().out)["chunks"] == 7
    assert cli.main(["eval"]) == 0
    assert json.loads(capsys.readouterr().out)["mrr"] == 0.5
